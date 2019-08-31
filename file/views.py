from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework import status
from rest_framework.decorators import api_view
from file.models import File
from file.serializers import FileSerializer
from docx import Document
from docx.shared import Inches

from subject.models import Subject
from user.models import User
from exercise.models import Exercise
from exam.models import Exam

from rest_framework.generics import ListCreateAPIView,RetrieveUpdateDestroyAPIView
import io;

# Create your views here.
class FileListCreateAPIView(ListCreateAPIView):
    queryset = File.objects.all()
    serializer_class = FileSerializer


class FileRetrieveUpdateDestroyAPIView(RetrieveUpdateDestroyAPIView):
    queryset = File.objects.all()
    serializer_class = FileSerializer
    lookup_field = 'id'


@csrf_exempt
@api_view(['POST', ])
def print_single_file(request):
    print('inside assign single Ip')
    if request.method == 'POST':

        # Load json data from body
        file_id = request.data['file_id']
        # date, exam, subject, full name of student + username+studentID+{exerciseName,submittedFile}

        file_instance = File.objects.get(pk=file_id)
        user_instance = User.objects.get(username=file_instance.user)
        exercise_instance = Exercise.objects.get(pk=file_instance.exercise.id)
        exam_instance = Exam.objects.get(pk=exercise_instance.exam.id)

        date = str(file_instance.date_created.strftime("%d/%m/%Y %H:%M:%S"))
        exam_name = exam_instance.name
        full_name = user_instance.first_name + user_instance.last_name
        username = user_instance.username
        exercise_name = exercise_instance.name
        file_name = file_instance.name

        header_text = date+" "+exam_name+" "+full_name+" "+username+" "+exercise_name+" "+file_name

        document = Document()
        header = document.sections[0].header
        head = header.paragraphs[0]
        head.text = header_text

        bf = file_instance.file_obj.read().decode('UTF-8')
        # bf.open(mode='rb')
        # lines = bf.readlines()
        print("bf.read()===", bf)
        document.add_paragraph(
            str(bf)
        )
        document.add_page_break()

        document.save('demo.docx')

        f = open("demo.docx", 'rb')
        from django.core.files import File as filexx

        file_obj = filexx(f)

        print("After convert ", type(file_obj), file_obj)
        new_file_name = "print_"+ file_name
        new_file = File(
            name=new_file_name,
            is_submitted=False,
            file_obj=file_obj,
            creator_id=user_instance.id,
            user=user_instance,
            exercise=None,
            subject_id=file_instance.subject.id
        )

        new_file.save()
        f.close()

        print(new_file)

        return JsonResponse({'msg': 'Successfully Assigned!', 'success': 1}, status=status.HTTP_200_OK)

    else:
        return JsonResponse({'msg': 'Method not allowed!', 'success': 0}, status=status.HTTP_405_METHOD_NOT_ALLOWED)



